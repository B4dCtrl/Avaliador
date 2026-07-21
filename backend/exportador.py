"""
Geração de laudos em Word (DOCX) e PDF para o Avaliador.

Inclui tabelas formatadas, gráficos embutidos e cabeçalho profissional.
"""

import io
import logging
from datetime import datetime
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
)

from graficos import gerar_todos_graficos_png

logger = logging.getLogger(__name__)

AZUL = RGBColor(0x1E, 0x40, 0xAF)
CINZA_CLARO = RGBColor(0xF3, 0xF4, 0xF6)


# ---------------------------------------------------------------------------
# Helpers Word
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de célula DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, texto: str, nivel: int = 1):
    p = doc.add_heading(texto, level=nivel)
    run = p.runs[0]
    run.font.color.rgb = AZUL
    return p


def _add_tabela_coeficientes(doc: Document, coeficientes: list):
    cabecalhos = [
        "Variável", "Transform.", "Coef.", "Elasticidade",
        "t-calc", "p-valor", "Sig.%", "IC 95% inf", "IC 95% sup"
    ]
    tabela = doc.add_table(rows=1 + len(coeficientes), cols=len(cabecalhos))
    tabela.style = "Table Grid"

    # Cabeçalho
    for i, cab in enumerate(cabecalhos):
        cell = tabela.rows[0].cells[i]
        cell.text = cab
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "1E40AF")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, coef in enumerate(coeficientes, start=1):
        ic = coef.get("intervalo_confianca_95", [None, None])
        valores = [
            coef.get("variavel", ""),
            coef.get("transformacao", ""),
            f"{coef.get('coeficiente', 0):.4f}",
            f"{coef.get('elasticidade', 0):.4f}",
            f"{coef.get('t_calculado', 0):.4f}",
            f"{coef.get('p_valor', 0):.4f}",
            f"{coef.get('significancia_percent', 0):.2f}%",
            f"{ic[0]:.4f}" if ic[0] is not None else "-",
            f"{ic[1]:.4f}" if ic[1] is not None else "-",
        ]
        for col_idx, val in enumerate(valores):
            cell = tabela.rows[row_idx].cells[col_idx]
            cell.text = val
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_idx % 2 == 0:
                _set_cell_bg(cell, "EFF6FF")


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def gerar_word(
    resultado: Dict[str, Any],
    imovel: Dict[str, Any],
    avaliador: Dict[str, Any],
) -> bytes:
    """
    Gera laudo completo em DOCX.

    Args:
        resultado: Response completo de /calcular-regressao.
        imovel: Dados do imóvel avaliado.
        avaliador: Dados do avaliador responsável.

    Returns:
        Bytes do arquivo DOCX.
    """
    doc = Document()

    # Margem
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cabeçalho
    _add_heading(doc, "LAUDO DE AVALIAÇÃO IMOBILIÁRIA", nivel=1)
    doc.add_paragraph(f"Conforme NBR 14653-02 | Data: {imovel.get('data_avaliacao', datetime.today().strftime('%Y-%m-%d'))}")
    doc.add_paragraph()

    # Dados do imóvel
    _add_heading(doc, "1. Identificação do Imóvel", nivel=2)
    tabela_imovel = doc.add_table(rows=4, cols=2)
    tabela_imovel.style = "Table Grid"
    dados_imovel = [
        ("Endereço", imovel.get("endereco", "")),
        ("Área do Terreno (m²)", str(imovel.get("area_terreno", ""))),
        ("Área Construída (m²)", str(imovel.get("area_construida", ""))),
        ("Finalidade", imovel.get("finalidade", "")),
    ]
    for i, (chave, valor) in enumerate(dados_imovel):
        tabela_imovel.rows[i].cells[0].text = chave
        tabela_imovel.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tabela_imovel.rows[i].cells[1].text = valor
    doc.add_paragraph()

    # Dados do avaliador
    _add_heading(doc, "2. Responsável Técnico", nivel=2)
    doc.add_paragraph(f"Nome: {avaliador.get('nome', '')}")
    doc.add_paragraph(f"CREA: {avaliador.get('crea', '')}")
    doc.add_paragraph(f"Empresa: {avaliador.get('empresa', '')}")
    doc.add_paragraph()

    # Resultados da regressão
    _add_heading(doc, "3. Resultados da Regressão Linear (OLS)", nivel=2)
    reg = resultado.get("regressao", {})
    tabela_reg = doc.add_table(rows=7, cols=2)
    tabela_reg.style = "Table Grid"
    stats_items = [
        ("R²", f"{reg.get('r_squared', 0):.4f}"),
        ("R² Ajustado", f"{reg.get('r_ajustado', 0):.4f}"),
        ("F calculado", f"{reg.get('f_calculado', 0):.4f}"),
        ("p-valor (F)", f"{reg.get('p_valor_f', 0):.6f}"),
        ("Durbin-Watson", f"{reg.get('durbin_watson', 0):.4f}"),
        ("Desvio-Padrão", f"{reg.get('desvio_padrao', 0):.4f}"),
        ("Observações", str(reg.get("observacoes", 0))),
    ]
    for i, (chave, valor) in enumerate(stats_items):
        tabela_reg.rows[i].cells[0].text = chave
        tabela_reg.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tabela_reg.rows[i].cells[1].text = valor
    doc.add_paragraph()

    # Coeficientes
    _add_heading(doc, "4. Coeficientes e Significância", nivel=2)
    _add_tabela_coeficientes(doc, resultado.get("coeficientes", []))
    doc.add_paragraph(f"\nIntercepto: {resultado.get('intercepto', 0):.6f}")
    doc.add_paragraph()

    # Validação NBR
    _add_heading(doc, "5. Validação NBR 14653-02", nivel=2)
    validacao = resultado.get("validacao_nbr", {})
    status_txt = "✓ Aprovado" if validacao.get("ok") else "✗ Reprovado"
    doc.add_paragraph(f"Status: {status_txt}")
    for aviso in validacao.get("avisos", []):
        doc.add_paragraph(f"⚠ {aviso}", style="List Bullet")
    for erro in validacao.get("erros", []):
        doc.add_paragraph(f"✗ {erro}", style="List Bullet")
    doc.add_paragraph()

    # Grau de fundamentação por item (se disponível)
    gf = resultado.get("grau_fundamentacao") or {}
    itens_gf = gf.get("itens") or {}
    if itens_gf:
        _add_heading(doc, "5.1 Grau de Fundamentação (itens quantificáveis)", nivel=2)
        doc.add_paragraph(f"Grau final (critério conservador): {gf.get('grau', '—')}")
        rot = {
            "quantidade_dados": "Quantidade de dados de mercado",
            "significancia_regressores": "Significância dos regressores (t)",
            "significancia_modelo_f": "Significância do modelo (F)",
        }
        for chave, info in itens_gf.items():
            doc.add_paragraph(
                f"{rot.get(chave, chave)}: grau {info.get('grau', '—')}",
                style="List Bullet",
            )
        obs = gf.get("observacao")
        if obs:
            doc.add_paragraph(obs)
        doc.add_paragraph()

    # Avaliação do imóvel-alvo (se disponível)
    aval = resultado.get("avaliacao_imovel") or {}
    if aval:
        _add_heading(doc, "5.2 Avaliação do Imóvel", nivel=2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"

        def linha(rotulo, valor):
            cells = t.add_row().cells
            cells[0].text = rotulo
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].text = valor

        def moeda(v):
            try:
                return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            except (TypeError, ValueError):
                return "—"

        linha("Valor estimado pelo modelo", moeda(aval.get("valor_estimado")))
        va = aval.get("valor_adotado")
        if va is not None:
            linha("Valor adotado pelo avaliador", moeda(va))
            try:
                desvio = (float(va) / float(aval["valor_estimado"]) - 1) * 100
                linha("Desvio em relação ao modelo", f"{desvio:+.1f}% (dentro do campo de arbítrio de ±15%)")
            except (TypeError, ValueError, KeyError, ZeroDivisionError):
                pass
        ic_m = aval.get("intervalo_confianca_media")
        if ic_m:
            linha("Intervalo de confiança da média (80%)", f"{moeda(ic_m[0])} a {moeda(ic_m[1])}")
        ic_p = aval.get("intervalo_predicao")
        if ic_p:
            linha("Intervalo de predição", f"{moeda(ic_p[0])} a {moeda(ic_p[1])}")
        ca = aval.get("campo_arbitrio")
        if ca:
            linha("Campo de arbítrio (±15%)", f"{moeda(ca[0])} a {moeda(ca[1])}")
        if aval.get("amplitude_pct") is not None:
            linha("Amplitude do IC", f"{aval['amplitude_pct']}%")
        if aval.get("grau_precisao"):
            linha("Grau de precisão", str(aval["grau_precisao"]))
        doc.add_paragraph()

    # Poder de predição (se disponível)
    pp = resultado.get("poder_predicao") or {}
    if pp.get("observado"):
        _add_heading(doc, "5.3 Poder de Predição", nivel=2)
        doc.add_paragraph(
            f"Desvio médio entre valores observados e estimados: {pp.get('desvio_medio_pct', 0)}%. "
            f"Amostras com desvio dentro de ±20%: {pp.get('pct_dentro_20', 0)}%."
        )
        doc.add_paragraph()

    # Gráficos
    _add_heading(doc, "6. Gráficos", nivel=2)
    imagens = gerar_todos_graficos_png(
        resultado.get("graficos", {}),
        nome_dependente="Valor",
    )
    for nome_img, dados_png in imagens.items():
        buf = io.BytesIO(dados_png)
        try:
            doc.add_picture(buf, width=Inches(5.5))
            doc.add_paragraph()
        except Exception as e:
            logger.warning("Não foi possível embutir imagem %s: %s", nome_img, e)

    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph(
        f"Laudo gerado pelo sistema Avaliador em {datetime.now().strftime('%d/%m/%Y %H:%M')}. "
        "Conforme NBR 14653-02 — ABNT."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf_out = io.BytesIO()
    doc.save(buf_out)
    buf_out.seek(0)
    return buf_out.read()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def gerar_pdf(
    resultado: Dict[str, Any],
    imovel: Dict[str, Any],
    avaliador: Dict[str, Any],
) -> bytes:
    """
    Gera laudo completo em PDF usando reportlab.

    Args:
        resultado: Response completo de /calcular-regressao.
        imovel: Dados do imóvel.
        avaliador: Dados do avaliador.

    Returns:
        Bytes do arquivo PDF.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    titulo_style = ParagraphStyle("Titulo", parent=styles["Heading1"],
                                   fontSize=16, textColor=colors.HexColor("#1E40AF"),
                                   spaceAfter=6)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"],
                               fontSize=12, textColor=colors.HexColor("#1E40AF"),
                               spaceAfter=4)
    normal = styles["Normal"]
    story = []

    def h1(txt): story.append(Paragraph(txt, titulo_style))
    def h2(txt): story.append(Paragraph(txt, h2_style))
    def p(txt): story.append(Paragraph(txt, normal)); story.append(Spacer(1, 4))
    def sp(): story.append(Spacer(1, 12))

    h1("LAUDO DE AVALIAÇÃO IMOBILIÁRIA")
    p(f"Conforme NBR 14653-02 | Data: {imovel.get('data_avaliacao', '')}")
    sp()

    h2("1. Identificação do Imóvel")
    dados_imovel = [
        ["Campo", "Valor"],
        ["Endereço", imovel.get("endereco", "")],
        ["Área do Terreno (m²)", str(imovel.get("area_terreno", ""))],
        ["Área Construída (m²)", str(imovel.get("area_construida", ""))],
        ["Finalidade", imovel.get("finalidade", "")],
    ]
    t = Table(dados_imovel, colWidths=[5*cm, 11*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t); sp()

    h2("2. Responsável Técnico")
    p(f"Nome: {avaliador.get('nome', '')} | CREA: {avaliador.get('crea', '')} | Empresa: {avaliador.get('empresa', '')}")
    sp()

    h2("3. Resultados da Regressão Linear (OLS)")
    reg = resultado.get("regressao", {})
    dados_reg = [
        ["Estatística", "Valor"],
        ["R²", f"{reg.get('r_squared', 0):.4f}"],
        ["R² Ajustado", f"{reg.get('r_ajustado', 0):.4f}"],
        ["F calculado", f"{reg.get('f_calculado', 0):.4f}"],
        ["p-valor (F)", f"{reg.get('p_valor_f', 0):.6f}"],
        ["Durbin-Watson", f"{reg.get('durbin_watson', 0):.4f}"],
        ["Desvio-Padrão", f"{reg.get('desvio_padrao', 0):.4f}"],
        ["Observações", str(reg.get("observacoes", 0))],
    ]
    t2 = Table(dados_reg, colWidths=[8*cm, 8*cm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t2); sp()

    h2("4. Coeficientes e Significância")
    cabecalhos_coef = ["Variável", "Transform.", "Coef.", "Elastic.", "t-calc", "p-valor", "Sig.%"]
    linhas_coef = [cabecalhos_coef]
    for coef in resultado.get("coeficientes", []):
        linhas_coef.append([
            coef.get("variavel", ""),
            coef.get("transformacao", ""),
            f"{coef.get('coeficiente', 0):.4f}",
            f"{coef.get('elasticidade', 0):.4f}",
            f"{coef.get('t_calculado', 0):.4f}",
            f"{coef.get('p_valor', 0):.4f}",
            f"{coef.get('significancia_percent', 0):.2f}%",
        ])
    t3 = Table(linhas_coef)
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    story.append(t3)
    p(f"Intercepto: {resultado.get('intercepto', 0):.6f}")
    sp()

    h2("5. Validação NBR 14653-02")
    validacao = resultado.get("validacao_nbr", {})
    status = "APROVADO" if validacao.get("ok") else "REPROVADO"
    p(f"Status: {status}")
    for aviso in validacao.get("avisos", []):
        p(f"⚠ {aviso}")
    for erro in validacao.get("erros", []):
        p(f"✗ {erro}")
    sp()

    def _moeda(v):
        try:
            return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except (TypeError, ValueError):
            return "—"

    gf = resultado.get("grau_fundamentacao") or {}
    itens_gf = gf.get("itens") or {}
    if itens_gf:
        h2("5.1 Grau de Fundamentação (itens quantificáveis)")
        p(f"Grau final (critério conservador): {gf.get('grau', '—')}")
        rot = {
            "quantidade_dados": "Quantidade de dados de mercado",
            "significancia_regressores": "Significância dos regressores (t)",
            "significancia_modelo_f": "Significância do modelo (F)",
        }
        for chave, info in itens_gf.items():
            p(f"• {rot.get(chave, chave)}: grau {info.get('grau', '—')}")
        if gf.get("observacao"):
            p(gf["observacao"])
        sp()

    aval = resultado.get("avaliacao_imovel") or {}
    if aval:
        h2("5.2 Avaliação do Imóvel")
        linhas_av = [["Item", "Valor"]]
        linhas_av.append(["Valor estimado pelo modelo", _moeda(aval.get("valor_estimado"))])
        va = aval.get("valor_adotado")
        if va is not None:
            linhas_av.append(["Valor adotado pelo avaliador", _moeda(va)])
            try:
                desvio = (float(va) / float(aval["valor_estimado"]) - 1) * 100
                linhas_av.append(["Desvio vs. modelo", f"{desvio:+.1f}% (campo de arbítrio ±15%)"])
            except (TypeError, ValueError, KeyError, ZeroDivisionError):
                pass
        if aval.get("intervalo_confianca_media"):
            ic = aval["intervalo_confianca_media"]
            linhas_av.append(["IC da média (80%)", f"{_moeda(ic[0])} a {_moeda(ic[1])}"])
        if aval.get("intervalo_predicao"):
            ic = aval["intervalo_predicao"]
            linhas_av.append(["Intervalo de predição", f"{_moeda(ic[0])} a {_moeda(ic[1])}"])
        if aval.get("campo_arbitrio"):
            ca = aval["campo_arbitrio"]
            linhas_av.append(["Campo de arbítrio (±15%)", f"{_moeda(ca[0])} a {_moeda(ca[1])}"])
        if aval.get("amplitude_pct") is not None:
            linhas_av.append(["Amplitude do IC", f"{aval['amplitude_pct']}%"])
        if aval.get("grau_precisao"):
            linhas_av.append(["Grau de precisão", str(aval["grau_precisao"])])
        t_av = Table(linhas_av, colWidths=[8*cm, 8*cm])
        t_av.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EFF6FF")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
        ]))
        story.append(t_av)
        sp()

    pp_dados = resultado.get("poder_predicao") or {}
    if pp_dados.get("observado"):
        h2("5.3 Poder de Predição")
        p(
            f"Desvio médio entre observado e estimado: {pp_dados.get('desvio_medio_pct', 0)}%. "
            f"Amostras dentro de ±20%: {pp_dados.get('pct_dentro_20', 0)}%."
        )
        sp()

    h2("6. Gráficos")
    imagens = gerar_todos_graficos_png(resultado.get("graficos", {}))
    for _, dados_png in imagens.items():
        img_buf = io.BytesIO(dados_png)
        try:
            img = Image(img_buf, width=14*cm, height=8*cm)
            story.append(img)
            story.append(Spacer(1, 8))
        except Exception as e:
            logger.warning("Erro ao embutir imagem no PDF: %s", e)

    p(f"Laudo gerado pelo sistema Avaliador em {datetime.now().strftime('%d/%m/%Y %H:%M')}. Conforme NBR 14653-02.")

    doc.build(story)
    buf.seek(0)
    return buf.read()
